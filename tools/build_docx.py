# -*- coding: utf-8 -*-
"""Preenche o modelo relatorioPOO.docx com o conteudo do relatorio.
Gera uma copia nova (Relatorio_LojaDeJogos.docx), mantendo o modelo intacto.
"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SRC = 'relatorioPOO.docx'
OUT = 'Relatorio_LojaDeJogos.docx'

doc = Document(SRC)
body = doc.element.body

# ---------- localizar cabecalhos Heading 1 (em ordem) ----------
HEAD_TXT = {
    'figuras': 'Lista de figuras',
    'intro': 'Introdução',
    'requisitos': 'Analise de requesitos',
    'estrutura': 'Estrutura do sistema',
    'trabalho': 'Trabalho realizado',
    'conclusoes': 'Conclusões e trabalho futuro',
    'referencias': 'Referências',
    'bibliografia': 'Bibliografia',
}
headings = {}
order = []
for p in doc.paragraphs:
    if p.style.name == 'Heading 1':
        for k, t in HEAD_TXT.items():
            if p.text.strip() == t:
                headings[k] = p
                order.append(p)
                break

def next_heading_elem(p):
    idx = order.index(p)
    return order[idx + 1]._p if idx + 1 < len(order) else None

def has_sectpr(el):
    return el.tag == qn('w:sectPr') or (el.tag == qn('w:p') and el.find(qn('w:pPr')) is not None
                                        and el.find(qn('w:pPr')).find(qn('w:sectPr')) is not None)

def clear_section(heading_p, keep=lambda el: False):
    stop = next_heading_elem(heading_p)
    el = heading_p._p.getnext()
    while el is not None and el is not stop:
        nxt = el.getnext()
        if not has_sectpr(el) and el.tag in (qn('w:p'), qn('w:tbl')) and not keep(el):
            el.getparent().remove(el)
        el = nxt

# ---------- helpers de insercao ----------
class Cursor:
    def __init__(self, heading_p):
        self.el = heading_p._p
    def place(self, element):
        self.el.addnext(element)
        self.el = element

def set_text(para, text):
    for r in list(para.runs):
        r.text = ''
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)

def add_p(cur, text='', mono=False, bold=False, italic=False, center=False, size=None):
    para = doc.add_paragraph()
    if text:
        run = para.add_run(text)
        if mono:
            run.font.name = 'Consolas'; run.font.size = Pt(size or 9)
        if size and not mono:
            run.font.size = Pt(size)
        run.bold = bold; run.italic = italic
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cur.place(para._p)
    return para

def add_head(cur, text, level=2):
    para = doc.add_paragraph(text)
    para.style = doc.styles[f'Heading {level}']
    cur.place(para._p)
    return para

def add_code(cur, text):
    para = doc.add_paragraph()
    for i, line in enumerate(text.split('\n')):
        if i > 0:
            para.add_run().add_break()
        r = para.add_run(line if line else ' ')
        r.font.name = 'Consolas'; r.font.size = Pt(8.5)
    cur.place(para._p)
    return para

def add_table(cur, rows):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = doc.styles['Table Grid']
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.cell(ri, ci)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(8.5)
            if ri == 0:
                run.bold = True
    cur.place(t._tbl)
    return t

def add_fig(cur, caption):
    add_p(cur, f'[{caption}]', italic=True, center=True)
    add_p(cur, '(inserir aqui o diagrama)', italic=True, center=True, size=9)

# =========================================================
# CONTEUDO
# =========================================================

# ----- Lista de figuras: remover apenas as duas Notas -----
def clear_figuras():
    stop = next_heading_elem(headings['figuras'])
    el = headings['figuras']._p.getnext()
    while el is not None and el is not stop:
        nxt = el.getnext()
        if el.tag == qn('w:p'):
            from docx.text.paragraph import Paragraph
            txt = Paragraph(el, doc).text.strip()
            if txt.startswith('Nota '):
                el.getparent().remove(el)
        el = nxt
clear_figuras()

# ----- Introdução -----
clear_section(headings['intro'])
cur = Cursor(headings['intro'])
add_p(cur, "Este relatório descreve a aplicação Loja de Jogos, desenvolvida em Java no âmbito da "
           "unidade curricular de Programação Orientada a Objectos. O dono de uma loja que vende "
           "produtos atuais, clássicos e vintage pretendia um programa que permitisse aos clientes "
           "consultar informação sobre os produtos (por título, diretor e outros detalhes) e que, ao "
           "mesmo tempo, processasse as vendas e mantivesse a contabilidade em dia.")
add_p(cur, "A aplicação é de consola e está dividida em duas zonas: uma zona de administrador, "
           "protegida por palavra-passe, onde se gere todo o catálogo (jogos, diretores, produtoras, "
           "produtos, localizações, clientes e promoções), se registam vendas e se consultam "
           "estatísticas; e uma zona de cliente, de acesso livre, dedicada a pesquisas sobre os jogos "
           "e produtos. Todos os dados persistem entre sessões, sendo guardados em ficheiros de "
           "objetos (.dat) e em ficheiros de texto (.txt).")

# ----- Analise de requisitos -----
set_text(headings['requisitos'], 'Análise de requisitos')
clear_section(headings['requisitos'])
cur = Cursor(headings['requisitos'])
add_p(cur, "Do enunciado foram levantados 13 requisitos (R1–R13). A tabela seguinte resume cada um "
           "e indica onde é cumprido no código.")
REQ = [
 ['Req', 'Descrição (resumo)', 'Onde é cumprido'],
 ['R1', 'A loja vende produtos (CD, DVD, Blu-ray, cartuchos, etc.).', 'model.Produto (formato)'],
 ['R2', 'Cada produto é composto por 1+ jogos; um jogo pode estar em vários produtos; versões diferentes são jogos diferentes.', 'Produto.jogos : List<Jogo>; cada Jogo é entidade própria'],
 ['R3', 'Cada jogo tem diretor, produtora e estilo.', 'Jogo (diretor, produtora, estilo)'],
 ['R4', 'Os produtos têm formato e plataforma (PS4, XBOX, PC, …).', 'Produto (formato, plataforma)'],
 ['R5', 'Localização (área/prateleira); produtos podem não existir na loja; guardar preço de custo e de venda.', 'Localizacao; Produto.localizacoes (pode ser vazia); precoCusto/precoVenda'],
 ['R6', 'Diretores: nome, data nasc., email, web, morada, morada clube de fãs, observações — em ficheiros de objetos.', 'model.Diretor; FileManager; diretores.dat'],
 ['R7', 'Produtoras em ficheiros de objetos; cliente vê diretores que trabalham numa produtora.', 'model.Produtora; produtoras.dat; LojaController.listarDiretoresDaProdutora'],
 ['R8', 'Pesquisar jogos por nome/diretor/estilo; produtos que contêm um jogo; produtos comprados pelos mesmos clientes.', 'LojaController.pesquisar*, produtosQueContemJogo, produtosCompradosPorClientesQueCompraram'],
 ['R9', 'Empregados inserem/atualizam; apagamentos permitidos para corrigir enganos.', 'CRUD em gestao.Gere* e ui.MenuAdmin'],
 ['R10', 'Registar vendas e diminuir stock; guardar empregado, item, quantidade, custo e data — em ficheiro de texto.', 'LojaController.registarVenda, Venda, vendas.txt'],
 ['R11', 'Clientes especiais (não anónimos) com registo das suas compras.', 'model.ClienteEspecial (historicoCompras)'],
 ['R12', 'Zonas admin/cliente; password em ficheiro de texto, encriptada após o 1.º login.', 'LoginManager, pass.txt, CryptoUtils (SHA-256)'],
 ['R13', 'Estatísticas de mais vendidos; promoções e saber se a venda teve promoção.', 'Estatísticas no LojaController; Promocao; Venda.promocao'],
]
add_table(cur, REQ)
add_p(cur, "Os requisitos agrupam-se em cinco famílias: gestão de entidades (R1–R7, R9), pesquisas "
           "(R7, R8), vendas e stock (R10, R11), segurança/acesso (R12) e estatísticas e promoções (R13).")

# ----- Estrutura do sistema -----
set_text(headings['estrutura'], 'Estrutura do sistema — Arquitetura em camadas')
clear_section(headings['estrutura'])
cur = Cursor(headings['estrutura'])
add_p(cur, "A aplicação está organizada segundo uma arquitetura em camadas, em que cada pacote tem "
           "uma responsabilidade única e as dependências fluem num só sentido:")
add_code(cur, "ui  →  controller  →  gestao  →  persistence\n        (model é transversal a todas as camadas)")
LAYERS = [
 ['Pacote', 'Camada', 'Responsabilidade'],
 ['ui', 'Apresentação', 'Menus de consola; lê o input e imprime resultados'],
 ['controller', 'Negócio / Serviço', 'LojaController — fachada com regras: venda, estatísticas, pesquisas, capacidade'],
 ['gestao', 'Serviço / Repositório', 'Classes Gere* — CRUD em memória e delegação da persistência'],
 ['persistence', 'Acesso a dados', 'FileManager<T> (objetos) e TextFileManager (texto)'],
 ['model', 'Domínio', 'Entidades de negócio (Jogo, Diretor, Produto, Venda, …)'],
 ['utils', 'Transversal', 'Input (validação) e CryptoUtils (SHA-256)'],
]
add_table(cur, LAYERS)
add_p(cur, "Optou-se por esta organização (e não pelo padrão MVC) porque descreve com rigor o que o "
           "código faz: não existe uma View a observar o Model, e a classe LojaController desempenha "
           "o papel de fachada de negócio e não de controlador no sentido do MVC. Ainda assim, o nome "
           "do pacote controller e o seu papel coordenador inspiram-se nessa ideia de mediação entre a "
           "apresentação e os dados.")
add_p(cur, "A separação tem vantagens claras: a camada ui desconhece como os dados são guardados; a "
           "classe FileManager<T> é genérica e serve qualquer tipo serializável, evitando repetição; e "
           "toda a validação de input está centralizada em utils.Input.")

# ----- Trabalho realizado -----
clear_section(headings['trabalho'])
cur = Cursor(headings['trabalho'])

add_head(cur, "Modelo de domínio", 2)
add_p(cur, "As entidades de negócio estão no pacote model e são todas Serializable (com "
           "serialVersionUID próprio), o que permite guardá-las em ficheiros de objetos.")
ENT = [
 ['Entidade', 'Atributos principais', 'Relações'],
 ['Jogo', 'nome, duração (min.), observações, estilo', '→ Diretor (0..1), → Produtora (0..1)'],
 ['Diretor', 'nome, data nasc., email, web, morada, morada clube de fãs, observações', '—'],
 ['Produtora', 'nome, morada, web, email, estilo principal, observações', '—'],
 ['Produto', 'formato, plataforma, preço custo, preço venda, stock, disponível', '→ Jogo (1..*), → Localizacao (0..*)'],
 ['Localizacao', 'tipo (ÁREA/PRATELEIRA), nome, capacidade, nº prateleira, descrição', '—'],
 ['Venda', 'empregado, quantidade, custo total, data/hora', '→ Produto (1), → Promocao (0..1), → ClienteEspecial (0..1)'],
 ['ClienteEspecial', 'nome, contacto, email, morada', '→ Venda (0..*) — histórico'],
 ['Promocao', 'nome, % desconto, data início, data fim, observações', '—'],
 ['Utilizador', 'username, hash da password, tipo (ADMIN/CLIENTE)', '—'],
]
add_table(cur, ENT)
add_p(cur, "Destacam-se duas regras do domínio: um produto é composto por um ou mais jogos e o mesmo "
           "jogo pode pertencer a vários produtos (R2); e a localização pode ser uma área ou uma "
           "prateleira (R5), resolvida com o enum Localizacao.Tipo. A Promocao sabe dizer se estaAtiva() "
           "numa data e como aplicarDesconto() a um preço.")

add_head(cur, "Persistência", 2)
add_p(cur, "Ficheiros de objetos (.dat): a classe genérica FileManager<T> serializa a lista completa "
           "de uma entidade (ObjectOutputStream) e lê-a (ObjectInputStream), devolvendo lista vazia se "
           "o ficheiro não existir. Cada entidade tem o seu ficheiro. Cumpre R6 e R7.")
add_p(cur, "Ficheiros de texto (.txt): TextFileManager lê/escreve linhas. Usado para o registo legível "
           "das vendas (vendas.txt, R10), a palavra-passe (pass.txt, R12) e a exportação das pesquisas "
           "do cliente. Cada classe Gere* carrega a lista no arranque e reescreve o ficheiro após cada "
           "modificação.")

add_head(cur, "Algoritmos principais", 2)

add_head(cur, "Autenticação e encriptação da palavra-passe (R12)", 3)
add_p(cur, "No primeiro arranque é criada uma password por defeito em texto simples. No primeiro login "
           "bem-sucedido, é substituída pelo seu resumo SHA-256 em Base64; daí em diante a verificação "
           "compara hashes. A deteção de 'ainda em texto simples' usa o comprimento (um hash SHA-256 em "
           "Base64 tem 44 caracteres).")
add_p(cur, "Pseudocódigo:", bold=True)
add_code(cur,
"""FUNÇÃO autenticarAdmin(password):
    SE NÃO existePassword(): RETORNAR falso
    conteudo ← ler(pass.txt)
    SE comprimento(conteudo) < 44 ENTÃO          // ainda em texto simples
        SE conteudo == password ENTÃO
            escrever(pass.txt, SHA256_Base64(password))   // encripta no 1.º login
            RETORNAR verdadeiro
        SENÃO RETORNAR falso
    SENÃO                                          // já encriptada
        RETORNAR SHA256_Base64(password) == conteudo""")
add_fig(cur, "Figura 3 — Fluxograma da autenticação e encriptação")

add_head(cur, "Registo de venda (R10, R13)", 3)
add_p(cur, "Valida o stock, aplica uma promoção ativa (se existir), diminui o stock, cria a venda, "
           "associa-a ao cliente especial (se houver) e regista-a no ficheiro de texto.")
add_p(cur, "Pseudocódigo:", bold=True)
add_code(cur,
"""FUNÇÃO registarVenda(empregado, produto, quantidade, cliente):
    SE produto.stock < quantidade: RETORNAR falso          // sem stock
    precoFinal ← produto.precoVenda × quantidade
    promoAtiva ← NULO
    PARA CADA p EM promoções:
        SE p.estaAtiva():
            promoAtiva ← p
            precoFinal ← p.aplicarDesconto(precoFinal)
            PARAR                                            // 1.ª promoção ativa
    produto.stock ← produto.stock − quantidade; guardar(produtos)
    venda ← nova Venda(empregado, produto, quantidade, precoFinal, agora, promoAtiva, cliente)
    adicionar(venda)
    SE cliente ≠ NULO: cliente.adicionarCompra(venda); guardar(clientes)
    escreverLinha(vendas.txt, formatar(venda))               // ficheiro de texto (R10)
    RETORNAR verdadeiro""")
add_fig(cur, "Figura 4 — Fluxograma do registo de venda")

add_head(cur, "Estatísticas — jogos mais vendidos (R13)", 3)
add_p(cur, "Percorre todas as vendas, conta as ocorrências de cada jogo e devolve a lista ordenada por "
           "contagem decrescente. O mesmo padrão serve as estatísticas por estilo, por diretor e por "
           "produtora, mudando apenas o critério de agrupamento.")
add_p(cur, "Pseudocódigo:", bold=True)
add_code(cur,
"""FUNÇÃO jogosMaisVendidos():
    contador ← mapa vazio <Jogo, inteiro>
    PARA CADA venda EM vendas:
        PARA CADA jogo EM venda.produto.jogos:
            contador[jogo] ← contador[jogo] + 1
    RETORNAR contador ORDENADO POR valor DESCENDENTE""")
add_p(cur, "Em Java usa-se a Stream API: flatMap para obter os jogos de cada venda e "
           "Collectors.groupingBy(jogo, counting()) para a contagem.", italic=True)
add_fig(cur, "Figura 5 — Fluxograma das estatísticas (jogos mais vendidos)")

add_head(cur, "Ponto de vista do utilizador", 2)
add_p(cur, "A aplicação arranca no menu principal com três opções: Zona Administrador (pede password), "
           "Zona Cliente e Sair (que guarda os dados).")
add_p(cur, "Zona de administrador — gere Diretores, Produtoras, Jogos, Produtos, Localizações, Clientes "
           "Especiais e Promoções (cada um com Listar / Adicionar / Editar / Remover), regista vendas, "
           "consulta estatísticas e altera a password.")
add_p(cur, "Zona de cliente — oito pesquisas (jogo por nome, por diretor, por estilo; produtos que "
           "contêm um jogo; outros jogos do mesmo diretor/estilo; produtos comprados pelos mesmos "
           "clientes; diretores de uma produtora) e permite exportar os resultados para um ficheiro de "
           "texto. O detalhe de utilização consta do manual do utilizador, em anexo.")

add_head(cur, "Ponto de vista do programador", 2)
add_p(cur, "Arquitetura em camadas: a ui desconhece a persistência; o LojaController é a fachada de "
           "negócio. Reutilização por genéricos: FileManager<T> serve qualquer entidade serializável e "
           "as classes Gere* partilham o mesmo padrão de CRUD. Validação centralizada em utils.Input "
           "(obrigatório/opcional, inteiros/decimais com intervalo, datas, sim/não), repetindo o pedido "
           "até a entrada ser válida. Identidade por chave de negócio: equals()/hashCode() em Produto "
           "(formato + plataforma + nomes dos jogos), ClienteEspecial (nome + email) e Localizacao (nome).")

add_head(cur, "Decisões e pressupostos", 2)
DEC = [
 "1. Identidade por chave de negócio. Como cada entidade é guardada no seu ficheiro, o 'mesmo' "
 "produto ou cliente surge como objetos diferentes após leitura. Sem equals()/hashCode() próprios, a "
 "pesquisa de R8 (produtos comprados pelos mesmos clientes) devolvia sempre vazio; definiu-se "
 "igualdade por chave de negócio.",
 "2. Promoção 'uma de cada vez'. Aplica-se a primeira promoção ativa da lista a qualquer produto; as "
 "promoções não estão associadas a produtos específicos. Simplificação assumida.",
 "3. Estatísticas por ocorrência. As contagens de mais vendidos contam uma ocorrência por venda, não "
 "ponderando a quantidade vendida (ver trabalho futuro).",
 "4. Segurança da password. SHA-256 em Base64 sem salt; o estado 'texto simples' é detetado pelo "
 "comprimento (< 44 caracteres). Adequado ao âmbito académico.",
 "5. Produtos sem existência física. A lista de localizações de um produto pode ficar vazia (R5).",
 "6. Versões de um jogo. Versões diferentes do mesmo jogo são entidades Jogo distintas (R2).",
 "7. Autenticação. A classe Utilizador está prevista no modelo, mas a autenticação do administrador é "
 "feita via pass.txt/LoginManager.",
]
for item in DEC:
    add_p(cur, item)

# ----- Conclusões -----
clear_section(headings['conclusoes'])
cur = Cursor(headings['conclusoes'])
add_p(cur, "O projeto cumpre os 13 requisitos do enunciado. Foi desenvolvida uma aplicação de consola "
           "em Java que permite gerir todo o catálogo da loja (jogos, diretores, produtoras, produtos, "
           "localizações, clientes especiais e promoções) com operações de listagem, inserção, edição e "
           "remoção; registar vendas com diminuição automática de stock e aplicação de promoções; "
           "consultar estatísticas; e disponibilizar ao cliente pesquisas com exportação para ficheiro "
           "de texto. O acesso de administrador é protegido por palavra-passe encriptada e os dados "
           "persistem entre sessões.")
add_p(cur, "Quanto ao como, adotou-se uma arquitetura em camadas que separa apresentação, negócio, "
           "gestão das entidades e persistência, apoiada na serialização de objetos (.dat) e em "
           "ficheiros de texto (.txt). A reutilização foi conseguida com uma classe de persistência "
           "genérica (FileManager<T>) e com a centralização da validação em utils.Input.")
add_p(cur, "As principais dificuldades foram a modelação das relações entre entidades e, sobretudo, a "
           "identidade dos objetos guardados em ficheiros distintos: o mesmo produto ou cliente surge "
           "como objetos diferentes após leitura. Resolveu-se com equals()/hashCode() por chave de "
           "negócio, sem o qual a pesquisa de R8 não funcionava entre sessões.")
add_p(cur, "Reconhecem-se limitações assumidas: a promoção é aplicada de forma global (a primeira "
           "ativa) e as estatísticas contam ocorrências por venda sem ponderar a quantidade.")
add_p(cur, "Como trabalho futuro: (i) unificar as oito classes Gere* numa única classe genérica; (ii) "
           "associar promoções a produtos ou categorias; (iii) ponderar as estatísticas pela quantidade "
           "vendida; (iv) permitir desabilitar produtos em vez de os remover; e (v) evoluir para uma "
           "interface gráfica.")

# ----- Referências -----
clear_section(headings['referencias'])
cur = Cursor(headings['referencias'])
for r in [
 "[1] Enunciado do Trabalho Prático — Programação Orientada a Objectos, CTeSP-TPSI, ESTG, 2025/26.",
 "[2] Oracle, Java Platform SE — API Specification: java.io.Serializable, ObjectOutputStream/ObjectInputStream. https://docs.oracle.com/en/java/javase/",
 "[3] Oracle, Java SE API: java.security.MessageDigest (SHA-256).",
 "[4] Oracle, Java SE API: java.util.stream (Stream API) e java.util.stream.Collectors.",
]:
    add_p(cur, r)

# ----- Bibliografia -----
clear_section(headings['bibliografia'])
cur = Cursor(headings['bibliografia'])
for b in [
 "BLOCH, Joshua. Effective Java, 3.ª ed. Addison-Wesley, 2018.",
 "Oracle. The Java Tutorials. https://docs.oracle.com/javase/tutorial/",
 "Material de apoio das aulas de Programação Orientada a Objectos (2025/26).",
]:
    add_p(cur, b)

doc.save(OUT)
print("Gerado:", OUT)
print("Tabelas:", len(doc.tables), "| Paragrafos:", len(doc.paragraphs))
